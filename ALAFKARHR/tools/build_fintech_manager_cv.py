from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate


OUT_DIR = Path("output/cv")
DOCX_PATH = OUT_DIR / "Bashir_Mohamed_Ali_FinTech_Software_Development_Manager_CV.docx"
PDF_PATH = OUT_DIR / "Bashir_Mohamed_Ali_FinTech_Software_Development_Manager_CV.pdf"

BLUE = RGBColor(31, 77, 120)
DARK = RGBColor(15, 23, 42)
MUTED = RGBColor(71, 85, 105)


def set_cell_text_noop():
    # Placeholder to keep helper section grouped; no tables are used for ATS readability.
    return None


def set_paragraph_spacing(paragraph, before=0, after=0, line=1.0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_run(run, size=10, bold=False, color=DARK, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def set_margins(section):
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)


def add_bottom_border(paragraph, color="CBD5E1", size="6", space="1"):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)


def add_text_paragraph(doc, text="", size=10, bold=False, color=DARK, align=None, before=0, after=2, italic=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    set_paragraph_spacing(p, before=before, after=after, line=1.03)
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, color=color, italic=italic)
    return p


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=2, line=1.0)
    r = p.add_run(text.upper())
    set_run(r, size=10.5, bold=True, color=BLUE)
    add_bottom_border(p)
    return p


def add_role(doc, title, org_date, summary=None):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=3, after=1, line=1.0)
    r = p.add_run(title)
    set_run(r, size=10.2, bold=True, color=DARK)
    r = p.add_run(f" | {org_date}")
    set_run(r, size=10.0, bold=False, color=MUTED)
    if summary:
        add_text_paragraph(doc, summary, size=9.5, color=DARK, before=0, after=1)


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    set_paragraph_spacing(p, before=0, after=1.2, line=1.0)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run(r, size=9.45, bold=True, color=DARK)
        r = p.add_run(text[len(bold_prefix):])
        set_run(r, size=9.45, color=DARK)
    else:
        r = p.add_run(text)
        set_run(r, size=9.45, color=DARK)
    return p


def add_compact_line(doc, label, value):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=1, line=1.0)
    r = p.add_run(label)
    set_run(r, size=9.4, bold=True, color=DARK)
    r = p.add_run(value)
    set_run(r, size=9.4, color=DARK)
    return p


def build_doc():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()

    section = doc.sections[0]
    set_margins(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.03

    for style_name in ["List Bullet", "List Bullet 2"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(9.45)
        style.paragraph_format.space_after = Pt(1.2)
        style.paragraph_format.line_spacing = 1.0

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(name, before=0, after=0, line=1.0)
    r = name.add_run("BASHIR MOHAMED ALI")
    set_run(r, size=18, bold=True, color=DARK)

    headline = add_text_paragraph(
        doc,
        "Software Development Manager | FinTech & Payments | .NET Microservices | Secure API Platforms",
        size=10.5,
        bold=True,
        color=BLUE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=1,
    )

    contact = add_text_paragraph(
        doc,
        "Saudi Arabia | +966507804458 | baryce@gmail.com | linkedin.com/in/bashirmali",
        size=9.3,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=4,
    )

    add_section_heading(doc, "Professional Summary")
    add_text_paragraph(
        doc,
        "Software development manager and technical delivery leader with 13+ years of software engineering experience across banking, payments, healthcare, and enterprise systems. Strong background in FinTech and digital banking delivery, secure REST APIs, API design, .NET/Blazor platforms, SQL Server, and cloud-native microservices architecture. Known for leading and mentoring developers, coordinating Agile delivery, improving performance and uptime, and translating business-critical financial workflows into reliable, secure, production-ready systems.",
        size=9.65,
        after=2,
    )

    add_section_heading(doc, "Core Leadership & Technical Strengths")
    add_compact_line(
        doc,
        "Leadership & Delivery: ",
        "team leadership, mentoring, Agile delivery, release planning, code reviews, vendor coordination, cross-functional coordination, technical documentation.",
    )
    add_compact_line(
        doc,
        "FinTech & Architecture: ",
        "payment solutions, digital banking workflows, secure REST APIs, API governance, microservices architecture, event-driven services, OWASP-aware secure development.",
    )
    add_compact_line(
        doc,
        "Technology Stack: ",
        ".NET 6/7/8, C#, ASP.NET Core, Blazor, EF Core, SQL Server, Docker, Kubernetes, OpenShift, Google Cloud, CI/CD, observability, Git, JWT, MFA, AES-256.",
    )

    add_section_heading(doc, "Professional Experience")
    add_role(doc, "Software Development Manager / Senior Full Stack Software Engineer", "Farmer's Commercial Bank, Sudan | Nov 2015 - Present")
    add_bullet(doc, "Led design and delivery of .NET Core and Blazor applications supporting critical digital banking and financial operations.")
    add_bullet(doc, "Designed and integrated secure REST APIs for banking systems, improving integration reliability and reducing data transfer errors by 40%.")
    add_bullet(doc, "Delivered custom payment solutions that improved transaction processing time by 20% and strengthened operational responsiveness.")
    add_bullet(doc, "Coordinated Agile delivery with business, infrastructure, QA, and support stakeholders to shorten release cycles by 25%.")
    add_bullet(doc, "Mentored junior developers through code reviews, technical guidance, debugging support, and onboarding, improving team productivity by 10%.")
    add_bullet(doc, "Implemented AES-256 encryption and MFA controls, reducing potential security vulnerabilities by 50% and supporting secure API development aligned with OWASP principles.")
    add_bullet(doc, "Improved system performance by 30%, cut server costs by 20%, and enhanced uptime by 15% through optimization, proactive maintenance, and scalable architecture.")
    add_bullet(doc, "Provided technical coordination with internal teams and external technology partners for banking platform changes, deployments, troubleshooting, and production support.")
    add_compact_line(doc, "Key technologies: ", ".NET 6/7/8, ASP.NET Core, Blazor, C#, REST APIs, SQL Server, EF Core, Docker, Kubernetes, CI/CD, Git, Microservices, Agile.")

    add_role(doc, "Cloud-Native Microservices Architecture Project", "Personal Project | 2025 - Present")
    add_bullet(doc, "Designed a production-style microservices architecture deployed on Kubernetes using MicroK8s, Docker, and clean architecture principles.")
    add_bullet(doc, "Built Product, Coupon, Order, Shopping Cart, and Email services with ASP.NET Core, EF Core, SQL Server, JWT authentication, and centralized API communication.")
    add_bullet(doc, "Implemented event-driven communication using RabbitMQ and MassTransit to support resilient service-to-service workflows.")
    add_bullet(doc, "Configured Kubernetes Deployments, Services, ConfigMaps, Secrets, and Ingress to demonstrate scalable cloud-native deployment patterns.")
    add_bullet(doc, "Aligned architecture with CI/CD readiness, API design discipline, SOLID principles, and maintainable service boundaries.")
    add_compact_line(doc, "Tech stack: ", ".NET 8, ASP.NET Core, RabbitMQ, MassTransit, Docker, Kubernetes, MicroK8s, REST APIs, JWT, EF Core, SQL Server.")

    doc.add_page_break()

    add_role(doc, "Full Stack Software Engineer", "48 Modal Hospital, Yemen | Jun 2011 - Apr 2015")
    add_bullet(doc, "Designed and developed hospital information systems that digitized operational workflows and improved data accuracy.")
    add_bullet(doc, "Integrated modules with existing infrastructure and provided ongoing support, troubleshooting, and enhancements for reliable operations.")

    add_role(doc, "Technical Support Director", "WonderTech Middle East, Yemen | May 2010 - May 2011")
    add_bullet(doc, "Managed maintenance teams delivering surveillance and security technology solutions for client environments.")
    add_bullet(doc, "Oversaw installation, integration, troubleshooting, and daily coordination for digital security systems.")

    add_section_heading(doc, "Education")
    add_text_paragraph(doc, "Bachelor of Computer Science - Hadhramout University of Science and Technology, Yemen | 2004 - 2008", size=9.6, after=2)

    add_section_heading(doc, "Certifications & Professional Training")
    training = [
        "Certificate in Digital Money - Digital Frontiers Institute (DFI), August 2022",
        "Preparing for Google Cloud Certification: Cloud DevOps Engineer Professional Certificate - Google Cloud/Coursera, January 2026",
        "IBM DevOps, Cloud, and Agile Foundations Specialization - IBM/Coursera, September 2024",
        "Application Security for Developers and DevOps Professionals - IBM/Coursera, October 2024",
        "Modular Monolith Architecture: .NET 8, CQRS, API Development, and Module Communication - Packt/Coursera, February-March 2026",
        "Selected DevOps coursework: Kubernetes/OpenShift, CI/CD, Monitoring & Observability, Microservices & Serverless, TDD/BDD, Git/GitHub",
    ]
    for item in training:
        add_bullet(doc, item)

    add_section_heading(doc, "Selected Achievements")
    achievements = [
        "Reduced transaction processing time by 20% through custom payment solution delivery.",
        "Reduced release time by 25% by improving CI/CD and Agile delivery practices.",
        "Strengthened data protection by 50% with AES-256 encryption and MFA controls.",
        "Reduced infrastructure costs by 20% and improved system performance by 30% through optimization.",
    ]
    for item in achievements:
        add_bullet(doc, item)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(footer, before=0, after=0, line=1.0)
    r = footer.add_run("Bashir Mohamed Ali - Software Development Manager CV")
    set_run(r, size=8, color=MUTED)

    doc.save(DOCX_PATH)
    print(DOCX_PATH.resolve())


def pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name="CVName", parent=styles["Normal"], fontName="Helvetica-Bold",
        fontSize=16, leading=18, alignment=TA_CENTER,
        textColor=colors.HexColor("#0F172A"), spaceAfter=1.5,
    ))
    styles.add(ParagraphStyle(
        name="CVHeadline", parent=styles["Normal"], fontName="Helvetica-Bold",
        fontSize=9.4, leading=11, alignment=TA_CENTER,
        textColor=colors.HexColor("#1F4D78"), spaceAfter=1,
    ))
    styles.add(ParagraphStyle(
        name="CVContact", parent=styles["Normal"], fontName="Helvetica",
        fontSize=8.2, leading=9.5, alignment=TA_CENTER,
        textColor=colors.HexColor("#475569"), spaceAfter=4,
    ))
    styles.add(ParagraphStyle(
        name="CVHeading", parent=styles["Normal"], fontName="Helvetica-Bold",
        fontSize=9.1, leading=10, textColor=colors.HexColor("#1F4D78"),
        spaceBefore=5, spaceAfter=2, borderColor=colors.HexColor("#CBD5E1"),
        borderWidth=0.4, borderPadding=1.5,
    ))
    styles.add(ParagraphStyle(
        name="CVBody", parent=styles["Normal"], fontName="Helvetica",
        fontSize=8.35, leading=9.45, alignment=TA_LEFT,
        textColor=colors.HexColor("#0F172A"), spaceAfter=1.8,
    ))
    styles.add(ParagraphStyle(
        name="CVSmall", parent=styles["CVBody"],
        fontSize=8.1, leading=9.15, spaceAfter=1.2,
    ))
    styles.add(ParagraphStyle(
        name="CVRole", parent=styles["Normal"], fontName="Helvetica-Bold",
        fontSize=8.9, leading=10, textColor=colors.HexColor("#0F172A"),
        spaceBefore=3, spaceAfter=1.5,
    ))
    styles.add(ParagraphStyle(
        name="CVBullet", parent=styles["CVSmall"],
        leftIndent=12, firstLineIndent=-6, bulletIndent=2, spaceAfter=1,
    ))
    return styles


def add_pdf_heading(story, styles, text):
    story.append(Paragraph(text.upper(), styles["CVHeading"]))


def add_pdf_bullet(story, styles, text):
    story.append(Paragraph(text, styles["CVBullet"], bulletText="-"))


def build_pdf():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    styles = pdf_styles()
    story = [
        Paragraph("BASHIR MOHAMED ALI", styles["CVName"]),
        Paragraph("Software Development Manager | FinTech & Payments | .NET Microservices | Secure API Platforms", styles["CVHeadline"]),
        Paragraph("Saudi Arabia | +966507804458 | baryce@gmail.com | linkedin.com/in/bashirmali", styles["CVContact"]),
    ]

    add_pdf_heading(story, styles, "Professional Summary")
    story.append(Paragraph(
        "Software development manager and technical delivery leader with 13+ years of software engineering experience across banking, payments, healthcare, and enterprise systems. Strong background in FinTech and digital banking delivery, secure REST APIs, API design, .NET/Blazor platforms, SQL Server, and cloud-native microservices architecture. Known for leading and mentoring developers, coordinating Agile delivery, improving performance and uptime, and translating business-critical financial workflows into reliable, secure, production-ready systems.",
        styles["CVBody"],
    ))

    add_pdf_heading(story, styles, "Core Leadership & Technical Strengths")
    strengths = [
        ("Leadership & Delivery", "team leadership, mentoring, Agile delivery, release planning, code reviews, vendor coordination, cross-functional coordination, technical documentation."),
        ("FinTech & Architecture", "payment solutions, digital banking workflows, secure REST APIs, API governance, microservices architecture, event-driven services, OWASP-aware secure development."),
        ("Technology Stack", ".NET 6/7/8, C#, ASP.NET Core, Blazor, EF Core, SQL Server, Docker, Kubernetes, OpenShift, Google Cloud, CI/CD, observability, Git, JWT, MFA, AES-256."),
    ]
    for label, value in strengths:
        story.append(Paragraph(f"<b>{label}:</b> {value}", styles["CVSmall"]))

    add_pdf_heading(story, styles, "Professional Experience")
    roles = [
        (
            "Software Development Manager / Senior Full Stack Software Engineer",
            "Farmer's Commercial Bank, Sudan | Nov 2015 - Present",
            [
                "Led design and delivery of .NET Core and Blazor applications supporting critical digital banking and financial operations.",
                "Designed and integrated secure REST APIs for banking systems, improving integration reliability and reducing data transfer errors by 40%.",
                "Delivered custom payment solutions that improved transaction processing time by 20% and strengthened operational responsiveness.",
                "Coordinated Agile delivery with business, infrastructure, QA, and support stakeholders to shorten release cycles by 25%.",
                "Mentored junior developers through code reviews, technical guidance, debugging support, and onboarding, improving team productivity by 10%.",
                "Implemented AES-256 encryption and MFA controls, reducing potential security vulnerabilities by 50% and supporting secure API development aligned with OWASP principles.",
                "Improved system performance by 30%, cut server costs by 20%, and enhanced uptime by 15% through optimization, proactive maintenance, and scalable architecture.",
                "Provided technical coordination with internal teams and external technology partners for banking platform changes, deployments, troubleshooting, and production support.",
            ],
            "Key technologies: .NET 6/7/8, ASP.NET Core, Blazor, C#, REST APIs, SQL Server, EF Core, Docker, Kubernetes, CI/CD, Git, Microservices, Agile.",
        ),
        (
            "Cloud-Native Microservices Architecture Project",
            "Personal Project | 2025 - Present",
            [
                "Designed a production-style microservices architecture deployed on Kubernetes using MicroK8s, Docker, and clean architecture principles.",
                "Built Product, Coupon, Order, Shopping Cart, and Email services with ASP.NET Core, EF Core, SQL Server, JWT authentication, and centralized API communication.",
                "Implemented event-driven communication using RabbitMQ and MassTransit to support resilient service-to-service workflows.",
                "Configured Kubernetes Deployments, Services, ConfigMaps, Secrets, and Ingress to demonstrate scalable cloud-native deployment patterns.",
                "Aligned architecture with CI/CD readiness, API design discipline, SOLID principles, and maintainable service boundaries.",
            ],
            "Tech stack: .NET 8, ASP.NET Core, RabbitMQ, MassTransit, Docker, Kubernetes, MicroK8s, REST APIs, JWT, EF Core, SQL Server.",
        ),
        (
            "Full Stack Software Engineer",
            "48 Modal Hospital, Yemen | Jun 2011 - Apr 2015",
            [
                "Designed and developed hospital information systems that digitized operational workflows and improved data accuracy.",
                "Integrated modules with existing infrastructure and provided ongoing support, troubleshooting, and enhancements for reliable operations.",
            ],
            None,
        ),
        (
            "Technical Support Director",
            "WonderTech Middle East, Yemen | May 2010 - May 2011",
            [
                "Managed maintenance teams delivering surveillance and security technology solutions for client environments.",
                "Oversaw installation, integration, troubleshooting, and daily coordination for digital security systems.",
            ],
            None,
        ),
    ]
    for title, meta, bullets, tech in roles:
        story.append(Paragraph(f"{title} <font color='#475569'>| {meta}</font>", styles["CVRole"]))
        for bullet in bullets:
            add_pdf_bullet(story, styles, bullet)
        if tech:
            story.append(Paragraph(f"<b>{tech.split(': ')[0]}:</b> {tech.split(': ', 1)[1]}", styles["CVSmall"]))
        if title == "Cloud-Native Microservices Architecture Project":
            story.append(PageBreak())

    add_pdf_heading(story, styles, "Education")
    story.append(Paragraph("Bachelor of Computer Science - Hadhramout University of Science and Technology, Yemen | 2004 - 2008", styles["CVSmall"]))

    add_pdf_heading(story, styles, "Certifications & Professional Training")
    training = [
        "Certificate in Digital Money - Digital Frontiers Institute (DFI), August 2022",
        "Preparing for Google Cloud Certification: Cloud DevOps Engineer Professional Certificate - Google Cloud/Coursera, January 2026",
        "IBM DevOps, Cloud, and Agile Foundations Specialization - IBM/Coursera, September 2024",
        "Application Security for Developers and DevOps Professionals - IBM/Coursera, October 2024",
        "Modular Monolith Architecture: .NET 8, CQRS, API Development, and Module Communication - Packt/Coursera, February-March 2026",
        "Selected DevOps coursework: Kubernetes/OpenShift, CI/CD, Monitoring & Observability, Microservices & Serverless, TDD/BDD, Git/GitHub",
    ]
    for item in training:
        add_pdf_bullet(story, styles, item)

    add_pdf_heading(story, styles, "Selected Achievements")
    achievements = [
        "Reduced transaction processing time by 20% through custom payment solution delivery.",
        "Reduced release time by 25% by improving CI/CD and Agile delivery practices.",
        "Strengthened data protection by 50% with AES-256 encryption and MFA controls.",
        "Reduced infrastructure costs by 20% and improved system performance by 30% through optimization.",
    ]
    for item in achievements:
        add_pdf_bullet(story, styles, item)

    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        rightMargin=0.58 * inch,
        leftMargin=0.58 * inch,
        topMargin=0.48 * inch,
        bottomMargin=0.45 * inch,
        title="Bashir Mohamed Ali - FinTech Software Development Manager CV",
        author="Bashir Mohamed Ali",
    )

    def footer(canvas, _doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 7)
        canvas.setFillColor(colors.HexColor("#64748B"))
        canvas.drawCentredString(letter[0] / 2, 0.27 * inch, "Bashir Mohamed Ali - Software Development Manager CV")
        canvas.restoreState()

    doc.build(story, onFirstPage=footer, onLaterPages=footer)
    print(PDF_PATH.resolve())


if __name__ == "__main__":
    build_doc()
    build_pdf()
